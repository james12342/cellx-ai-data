from pathlib import Path
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path("outputs/business_plan_20260831")
LOGO = Path("rdp-marketing-site/assets/logo.png")
VIS = OUT / "visuals"
DOCX = OUT / "CellX_RDP_Business_Plan_Competition_English_Presenters.docx"


BLUE = "1F4FD8"
NAVY = "14213D"
CYAN = "10B8D9"
LIGHT = "EAF3FF"
MUTED = "50627A"


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="111827"):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    r.font.size = Pt(9.2)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_para(doc, text, size=10.4, bold=False, color="1F2937", space_after=5, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.08
    r = p.add_run(text)
    r.font.name = "Aptos"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.color.rgb = RGBColor.from_string(color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = "Aptos Display"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
    r.font.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
    r.font.size = Pt(22 if level == 1 else 15)
    return p


def add_visual(doc, filename, width=7.05):
    image = VIS / filename
    if image.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(image), width=Inches(width))
        p.paragraph_format.space_after = Pt(8)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.style = doc.styles["List Bullet"]
        p.paragraph_format.space_after = Pt(3)
        r = p.add_run(item)
        r.font.name = "Aptos"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
        r.font.size = Pt(9.8)
        r.font.color.rgb = RGBColor.from_string("1F2937")


def add_metric_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, label in enumerate(["Metric", "CellX Position", "Why It Matters"]):
        set_cell_text(hdr[i], label, True, "FFFFFF")
        shade_cell(hdr[i], NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            if i == 0:
                shade_cell(cells[i], LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, label in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], label, True, "FFFFFF")
        shade_cell(table.rows[0].cells[i], NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
            if i == 0:
                shade_cell(cells[i], LIGHT)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


pages = [
    {
        "title": "1. Executive Summary",
        "body": [
            "CellX Rapid Development Platform (CellX RDP) is a software-defined business platform for building data-centric applications, workflow automations, and AI-assisted operational tools without rewriting a customer's core system. Order management is only one scenario. The broader vision covers marketing, sales, commerce, inventory, supply chain, customer support, finance, administration, and any industry-specific process that can be modeled as data plus workflow.",
            "The product thesis is software-defined software: business applications should be reconfigured through metadata, workflow, permissions, connectors, and AI skills instead of being rewritten as one-off custom code. In the AI-agent era, this becomes a major demand because companies want one person or a very small team to operate more processes, launch more campaigns, manage more customer interactions, and reduce repetitive work without adding headcount.",
            "The opportunity is clear: companies need new internal applications faster, but many are blocked by legacy systems, custom database schemas, integration complexity, and AI token-cost concerns. CellX solves this by turning database tables, field metadata, permissions, import/export rules, and extension points into configurable business apps. Its AI Workflow Skill Designer adds a drag-and-drop layer for third-party systems, AI models, manual ChatGPT handoffs, and direct CellX database actions."
        ],
        "bullets": [
            "Product: software-defined business apps plus an AI workflow and agent-skill sidecar.",
            "Vision: software-defined software for flexible, extensible business processes.",
            "AI-era operating model: one skilled operator can coordinate marketing, sales, operations, support, data workflows, and human approval.",
            "Users: founders, one-person companies, SMB teams, e-commerce operators, manufacturers, education administrators, agencies, and internal IT teams.",
            "Business model: tiered SaaS and private Lightsail/cloud deployments from $99 to $899+ per month.",
            "Competition goal: show a working prototype, a clear market, and a path to paid pilots."
        ],
    },
    {
        "title": "2. The Problem",
        "body": [
            "Most companies face the same tension: business users need new screens, tables, approvals, imports, exports, integrations, AI workflows, and reporting quickly, while the engineering team must protect the stability of the existing backend. The result is a backlog of small but urgent applications across the whole company: product promotion, AI video content, influencer outreach, social media posting, lead capture, CRM follow-up, order management, inventory, shipment tracking, supplier research, customer support, finance approvals, and administrative dashboards.",
            "Traditional custom development is flexible but slow. Generic low-code tools are fast but often disconnected from the company's real database, permission model, and deployment environment. AI tools can accelerate reasoning, but they frequently sit outside the actual workflow and raise concerns about data handling, governance, and API token costs."
        ],
        "bullets": [
            "Operational teams need applications that match their real database and process rules.",
            "Developers need extension patterns that do not break the core backend.",
            "Companies want AI assistance, but not every AI step should consume platform-paid tokens.",
            "Lean teams and one-person companies need AI agents to execute repeatable business workflows across promotion, sales, delivery, support, and finance, not just answer questions.",
            "Integrations with Amazon, FedEx, UPS, Stripe, PayPal, Outlook, Google Docs, and internal databases are increasingly expected."
        ],
    },
    {
        "title": "3. The Solution",
        "body": [
            "CellX RDP addresses the gap with a configuration-first architecture. Tables and page metadata become business screens. Permissions, import/export, validation, and workflow hooks can be configured rather than hand-coded each time. The platform preserves the existing backend while enabling extension modules, sidecar APIs, and workflow pages to be added around it.",
            "This makes CellX a software-defined software layer. The user defines what the business needs - fields, forms, approvals, data operations, third-party connectors, AI decisions, marketing actions, sales follow-up, fulfillment steps, and support logic - and CellX turns those definitions into a working operational interface.",
            "The new AI Workflow Designer extends the platform beyond CRUD generation. Users can drag nodes, connect steps, test third-party credentials, show connection status, import/export workflow templates, and use manual ChatGPT or web handoff when they want AI reasoning without burning the platform owner's token budget."
        ],
        "bullets": [
            "Generate production-ready admin pages from database structure and metadata.",
            "Add workflow logic, conditions, loops, joins, approvals, and integration nodes visually.",
            "Activate CellX database tables as workflow nodes for query, create, update, delete, export, and audit actions.",
            "Package full business processes such as product promotion, lead generation, order fulfillment, inventory alerts, supplier analysis, and customer support.",
            "Support both API-key AI automation and manual ChatGPT handoff."
        ],
    },
    {
        "title": "4. Product Modules",
        "body": [
            "CellX is positioned as a modular platform rather than a single page builder or an order-management product. The current prototype includes a marketing site, an admin backend, a workflow designer, a node library, a sidecar API concept, and import/export templates. These pieces form a general-purpose business operating layer that can be adapted to many industries and workflows."
        ],
        "table": [
            ("Runtime Page Builder", "Build CRUD pages, search forms, tables, and actions from metadata for any business object.", "Cuts repetitive internal app development work."),
            ("CellX Backend", "Existing deployed application with user, role, menu, notice, generated table, order, and configurable business structures.", "Provides a real operational system instead of a demo-only shell."),
            ("AI Workflow Designer", "Drag-and-drop workflow canvas with node library, connectors, import/export templates, and test status.", "Turns integrations and AI steps into reusable skills."),
            ("Marketing Site", "English front page with pricing, payment page, trial/demo links, and AI workflow promotion.", "Creates a public sales entry point."),
            ("Sidecar Extension API", "Independent /ext-api layer for new integrations and workflow services.", "Adds capability without modifying the original JAR source.")
        ],
    },
    {
        "title": "5. AI Workflow Skill Designer",
        "body": [
            "The Workflow Designer is a visual automation surface designed for business operators and technical implementers. It includes draggable nodes, a grid canvas, node headers with category colors and icons, editable properties, connection testing, status indicators, and right-click removal for unused nodes.",
            "Its node library is organized into practical categories: marketing, social media, content generation, influencer outreach, triggers, logic and control, office documents, email/calendar/chat, commerce and CRM, shipping/payment/accounting, inventory and supply chain, customer support, data/devops/internal systems, CellX database, AI models, and agent skills."
        ],
        "bullets": [
            "Marketing workflow: trend research -> AI video/script generation -> influencer discovery -> TikTok/Facebook distribution -> lead capture -> CRM follow-up.",
            "Commerce workflow: Amazon bestseller research -> SKU lookup -> filter records -> export Excel -> manual ChatGPT supplier analysis -> email results.",
            "Operations workflow: inventory threshold -> supplier quote request -> approval -> purchase order -> shipment tracking -> support notification.",
            "AI modes: platform API key, user's own API key, and manual web handoff.",
            "Integration readiness: properties panels collect API keys, OAuth settings, account credentials, sandbox/live modes, scopes, and endpoint details.",
            "Governance fit: status badges and logs make workflow behavior visible to admins."
        ],
    },
    {
        "title": "6. CellX Data Activation",
        "body": [
            "A key differentiator is the CellX Database node library. Rather than treating workflow as a separate automation island, CellX exposes internal tables as controlled workflow resources. This allows users to query campaigns, leads, customers, orders, inventory, suppliers, tickets, invoices, approvals, and audit events or trigger follow-up tasks from data already stored in the platform.",
            "This approach can be implemented with safeguards: allow-listed tables, parameterized queries, read/write permission policies, audit logs, and staged review before destructive actions. That creates a practical route for expanding a deployed backend even when the original source code is unavailable."
        ],
        "bullets": [
            "Query table, create record, update record, delete record, export Excel, audit log, and custom SQL-template actions.",
            "Start with non-destructive read/export nodes, then graduate to controlled write actions.",
            "Keep the extension sidecar independent from the original backend package.",
            "Use database metadata to make workflow actions self-documenting."
        ],
    },
    {
        "title": "7. Target Customers",
        "body": [
            "CellX is strongest where teams have frequent internal application needs, structured data, and integration pressure. The product is designed to be cross-industry: e-commerce, consumer brands, agencies, manufacturing, education, services, logistics, healthcare administration, real estate operations, and SMB IT teams that cannot afford long custom development cycles.",
            "The product also fits agencies and implementation partners who repeatedly build admin portals for clients. For those teams, CellX can become a repeatable delivery engine with private deployment options."
        ],
        "table": [
            ("Retail / E-commerce", "Promotion, social media, listings, orders, SKUs, carrier labels, returns, supplier research.", "High workflow frequency and clear ROI."),
            ("Manufacturing / Supply Chain", "Inventory, equipment, maintenance, procurement, supplier scoring, approvals, event logs.", "Structured operational data needs fast custom screens."),
            ("Education Admin", "Class, student, enrollment, notification, and reporting tables.", "Many simple internal workflows and role-based access needs."),
            ("Professional Services / Agencies", "Client portals, campaign operations, admin backends, import/export tools.", "Repeatable platform creates margin leverage.")
        ],
    },
    {
        "title": "8. Market Opportunity",
        "body": [
            "The market is moving toward low-code platforms that are also AI-aware, integration-ready, and governed. Gartner projects the low-code development technologies market to reach $58.2 billion by 2029 with a 14.1% CAGR, citing agentic AI, citizen development, and operational excellence as adoption drivers. Forrester estimated the low-code and digital process automation market at $13.2 billion by the end of 2023 and approximately $30 billion by 2028, with an AI-fueled scenario approaching $50 billion.",
            "AI adoption creates an additional demand layer. Deloitte's 2026 State of AI reporting shows expanded sanctioned AI access and strong interest in customized agents, while McKinsey highlights that scaling agentic AI requires governed, reusable assets that can coordinate tools, data, and transactional systems. CellX is designed for this intersection: low-code business apps plus governed workflow and AI execution across the full company lifecycle."
        ],
        "table": [
            ("Low-code growth", "Gartner: $58.2B projected by 2029.", "Large and expanding budget category."),
            ("AI-infused development", "Forrester: AI could push low-code/DPA toward $50B by 2028.", "Supports the timing of CellX's AI workflow layer."),
            ("One-person AI company", "Founders and lean teams need agent-enabled operating systems for marketing, sales, delivery, and support.", "Creates demand for tools that compress headcount and process cost."),
            ("Agent governance gap", "Deloitte and McKinsey emphasize governance, reuse, and scale challenges.", "CellX can compete on controlled workflow design, not raw chatbot novelty.")
        ],
    },
    {
        "title": "9. Competitive Landscape",
        "body": [
            "For an investor audience, the right competitive set is no longer only low-code tools. CellX should be compared with the broader AI workflow and enterprise agent ecosystem: Palantir AIP, Salesforce Agentforce, ServiceNow AI workflows, UiPath agentic automation, C3.ai enterprise AI applications, Microsoft Copilot and Power Platform, plus private AI-agent builders. These companies show that buyers want AI connected to data, workflow, governance, and execution.",
            "The public-market signal is mixed and useful. Palantir has been rewarded because investors see rapid growth tied to operational data and AI execution. UiPath and ServiceNow show that automation and workflow platforms can remain valuable when they absorb AI into existing process systems. Salesforce shows that agents can become an upsell layer inside a large application suite. C3.ai shows the risk: AI positioning alone is not enough if revenue growth weakens. CellX should use this pattern to position itself as a focused, SMB-friendly AI workflow and software-defined operations platform."
        ],
        "custom_table": {
            "headers": ["Public Comp", "AI / Workflow Angle", "Recent Revenue Signal", "Approx. 2-Year Stock Signal"],
            "rows": [
                ("Palantir (PLTR)", "AIP, data ontology, operational AI execution.", "2025 revenue $4.48B, +56%; TTM to Q2 2026 about $6.16B.", "Aug 2024 to Aug 2026: about +492%."),
                ("Salesforce (CRM)", "Agentforce, Data Cloud, CRM workflow agents.", "FY2026 revenue $41.5B, +10%; Agentforce/Data 360 ARR near $3.9B per recent reporting.", "Aug 2024 to Aug 2026: about +3%."),
                ("ServiceNow (NOW)", "Enterprise workflow platform with AI assistants and agents.", "2025 revenue $13.28B, +21%; AI workflow ACV reported above $1B in market coverage.", "Aug 2024 to Aug 2026: about -15% split-adjusted."),
                ("UiPath (PATH)", "RPA, automation, and agentic workflow execution.", "FY2026 revenue $1.61B, +13%; ARR $1.85B, +11%.", "Aug 2024 to Aug 2026: about +41%."),
                ("C3.ai (AI)", "Enterprise AI applications and agentic AI platform.", "FY2026 revenue $250M, -36%.", "Aug 2024 to Aug 2026: about -55%."),
                ("Microsoft (MSFT)", "Copilot, Azure AI, Power Platform, enterprise cloud.", "FY2026 revenue $331.8B, +18%; Microsoft Cloud $214.4B, +27%.", "Aug 2024 to Aug 2026: about +25%.")
            ],
        },
    },
    {
        "title": "10. Competitor Analysis",
        "body": [
            "The strongest investor framing is that CellX is not another generic AI chatbot or agent playground. It is closer to a lightweight operating layer for AI-enabled business workflows. Palantir is the clearest strategic reference because it connects data, workflow, governance, and operational decisions. However, Palantir primarily sells to larger enterprise and government buyers. CellX can use a smaller-company wedge: e-commerce operators, agencies, one-person companies, and SMB teams that need practical workflows without a heavy enterprise platform.",
            "Salesforce and ServiceNow show that agents become valuable when they live inside systems of record. UiPath shows that automation buyers still pay for execution, monitoring, and governance. Microsoft shows that AI distribution can be embedded into existing productivity and cloud workflows. CellX's opening is to package those ideas into a narrower, faster, lower-cost product: generated business apps, database workflow nodes, AI handoff options, templates, and private deployment."
        ],
        "custom_table": {
            "headers": ["Competitor", "What They Do Well", "CellX Opening", "Investor Read"],
            "rows": [
                ("Palantir AIP", "Operational AI, data ontology, enterprise/government deployment.", "Lighter SMB version of governed AI operations.", "Market rewards AI tied to data and execution."),
                ("Salesforce Agentforce", "CRM agents and Data Cloud inside a large app suite.", "Operate beyond CRM: marketing, commerce, supply chain, admin.", "Agents monetize when tied to workflow context."),
                ("ServiceNow AI", "IT, service, and enterprise workflow system of record.", "Lower-cost workflow layer for smaller teams and private deployments.", "Workflow incumbents are adding AI, validating the category."),
                ("UiPath agents", "RPA execution, automation governance, enterprise process automation.", "Generated apps plus database workflow before deep RPA complexity.", "Execution and governance matter more than chat UI."),
                ("C3.ai", "Enterprise AI applications and industry AI story.", "Use C3.ai as a cautionary comp: prove growth and retention early.", "AI label alone does not protect valuation.")
            ],
        },
    },
    {
        "title": "11. SWOT Analysis",
        "body": [
            "The SWOT view highlights a practical path for investor readiness. CellX already has a working prototype, a public site, a backend-adjacent architecture, and a strong AI-era narrative. The main weakness is that the product still needs customer proof, deeper connectors, and repeatable onboarding. The opportunity is large because low-code, agentic AI, workflow automation, and lean-company operations are converging. The threats are real: large platforms can bundle features, buyers will demand security, and the automation market is noisy.",
            "The strategy should therefore convert product progress into proof points: paid pilots, measurable time savings, reusable workflow templates, partner delivery, and documented security controls."
        ],
        "custom_table": {
            "headers": ["Strengths", "Weaknesses", "Opportunities", "Threats"],
            "rows": [
                ("Working prototype and deployed public/backend surfaces.", "Early brand awareness and limited customer references.", "Low-code and AI-agent demand is expanding.", "Large platforms can bundle adjacent features."),
                ("Backend-adjacent database activation and workflow nodes.", "Connector library still needs depth and reliability.", "SMBs want automation without hiring large teams.", "Security and compliance expectations can slow pilots."),
                ("AI cost modes: platform API, BYOK, manual handoff.", "Founder-led implementation may constrain early scale.", "Templates can turn services work into software margin.", "Crowded automation market makes positioning critical.")
            ],
        },
    },
    {
        "title": "12. Business Model and Pricing",
        "body": [
            "The current marketing site already presents a tiered structure that can be refined during customer discovery. The model balances self-service entry with higher-value private deployments and implementation support. The business should not rely on subscription revenue alone in the first year. Early revenue should come from paid pilots, setup packages, custom workflow templates, and connector configuration. As the product matures, the same work should become reusable software assets with higher gross margin.",
            "The most important pricing principle is to separate platform value from pass-through costs. CellX can charge for generated applications, workflow governance, private deployment, templates, and support, while AI model usage can be platform-managed, customer-owned through BYOK, or kept outside the platform through manual ChatGPT/web handoff."
        ],
        "table": [
            ("Trial", "$0 / month", "14-day demo, sample pages, workflow preview, and guided evaluation."),
            ("Starter", "$99 / month", "1 production environment, up to 20 generated pages, standard import/export, email support."),
            ("Growth", "$299 / month", "3 environments, unlimited generated pages, workflow extension support, priority updates."),
            ("Enterprise", "$899+ / month", "Private Lightsail/cloud deployment, source-level extension guidance, onboarding, governance review.")
        ],
        "bullets": [
            "Additional revenue can come from implementation packages, custom integrations, workflow templates, premium agent nodes, and industry solution packs.",
            "Template packs can cover AI marketing automation, social distribution, influencer outreach, sales follow-up, order fulfillment, inventory planning, supplier sourcing, support automation, and finance workflows.",
            "AI token costs can be passed through, BYOK-based, or avoided through manual ChatGPT/web handoff workflows.",
            "Private deployment keeps the platform attractive to companies with data-control requirements."
        ],
    },
    {
        "title": "13. Go-To-Market Strategy",
        "body": [
            "The first go-to-market motion should be founder-led and proof-driven. Instead of selling an abstract platform, CellX should demonstrate complete one-person-company workflows: generate product promotion videos, invite influencers, push content to TikTok/Facebook and other social platforms, capture leads, follow up through CRM/email, manage orders, check inventory, coordinate suppliers, handle shipping, support customers, and export results to Excel for AI analysis.",
            "The best beachhead is not every industry at once. CellX should start with customers who already feel the pain weekly: e-commerce operators, small agencies, brand owners, and SMB operations teams. These buyers have visible repetitive work, existing spreadsheets or databases, and a clear desire to reduce manual labor. The launch motion should produce short public demos, direct founder outreach, pilot workshops, and case studies showing time saved, workflows automated, and manual handoffs reduced."
        ],
        "bullets": [
            "Lead message: software-defined business processes for AI-era cost reduction and efficiency gains.",
            "Show the one-person-company story: one operator uses CellX to run promotion, sales, fulfillment, supplier analysis, and customer follow-up.",
            "Phase 1: paid pilots with e-commerce, marketing, agency, and SMB operations teams.",
            "Phase 2: package workflow templates for AI video promotion, influencer outreach, social posting, Amazon research, FedEx/UPS shipping, Stripe/PayPal payment follow-up, and CRM updates.",
            "Phase 3: partner with agencies that build marketing operations and internal systems for small businesses.",
            "Phase 4: publish a template marketplace and integration library."
        ],
    },
    {
        "title": "14. Market Launch Plan",
        "body": [
            "CellX should launch through a narrow, repeatable workflow wedge. The recommended wedge is AI marketing-to-sales operations for small teams: research a product or niche, create campaign assets, identify influencers or channels, publish or prepare posts, capture leads, update CRM records, and trigger follow-up. This is easy to demonstrate, strongly connected to revenue, and broad enough to expand into order, support, and supplier workflows.",
            "The launch plan should combine direct selling with productized content. A founder-led sales motion can target 50-100 qualified pilot prospects per month using LinkedIn, e-commerce communities, local SMB networks, agency relationships, and existing CellX site traffic. The public site should focus on demo booking, workflow templates, and before/after operational examples instead of generic platform language."
        ],
        "custom_table": {
            "headers": ["Channel", "Execution", "Success Metric", "Purpose"],
            "rows": [
                ("Founder outreach", "Direct email, LinkedIn, local business contacts, agency introductions.", "10-15 discovery calls per month.", "Find urgent pilot problems and refine ICP."),
                ("Demo content", "Short workflow videos for AI marketing, Amazon research, inventory alerts, support routing.", "5 qualified demo requests per month.", "Make the product understandable quickly."),
                ("Pilot workshops", "Paid setup session that maps one workflow and launches a prototype.", "3 paid pilots per quarter.", "Convert interest into revenue and proof."),
                ("Agency partners", "Offer private deployment and reusable client workflow templates.", "2 active partners in Year 1.", "Lower CAC and create implementation leverage.")
            ],
        },
    },
    {
        "title": "15. Technology and Architecture",
        "body": [
            "The deployed backend uses a Java-oriented enterprise stack, while the extension layer can be independently deployed behind Nginx. This allows CellX to grow even when the original backend source is unavailable. The workflow UI can be hosted under /workflow/ and communicate with a sidecar /ext-api layer for integrations, database operations, and AI connectors.",
            "The architecture prioritizes separation: the original system remains stable, while new workflow capabilities, integration adapters, and AI skill pages can be added beside it. This is the technical foundation for flexible business-process extension: a company can keep its core software stable while allowing AI agents and workflow nodes to operate through governed sidecar services."
        ],
        "table": [
            ("Frontend", "Vue3/Vite style UI plus standalone workflow designer and marketing site.", "Fast iteration and clean deployment."),
            ("Backend", "Java/Spring Boot style CellX backend with MySQL and generated-page metadata.", "Mature enterprise patterns."),
            ("Workflow Sidecar", "Python/FastAPI or equivalent /ext-api service.", "Independent extension without rebuilding core JAR."),
            ("Infrastructure", "AWS Lightsail, Nginx, Cloudflare DNS, app.cellaidata.com routing.", "Affordable deployment with upgrade path.")
        ],
    },
    {
        "title": "16. Security, Governance, and AI Cost Strategy",
        "body": [
            "For competition and pilot readiness, security should be presented as a core design principle. CellX should avoid giving arbitrary SQL or unrestricted API access to workflow users. Instead, it should use allow-listed database actions, least-privilege credentials, encrypted secrets, audit logs, environment separation, and clear manual approval gates.",
            "The AI strategy is especially important. Some users want automation through OpenAI, Anthropic, Gemini, DeepSeek, Mistral, Llama, Perplexity, or local Ollama models. Others want to log into ChatGPT manually and paste results back to avoid platform-paid token usage. CellX can support all three modes: platform-managed API, customer-provided API key, and human web handoff."
        ],
        "bullets": [
            "Secrets: store API keys, access keys, and passwords encrypted and scoped by workspace.",
            "Database: use parameterized queries and allow-listed tables/actions.",
            "AI: label token-cost ownership clearly for every model node.",
            "Efficiency: let AI agents reduce repetitive labor while keeping humans in control of sensitive decisions.",
            "Audit: record workflow runs, node inputs, outputs, status, and operator approvals."
        ],
    },
    {
        "title": "17. Roadmap",
        "body": [
            "The roadmap should move from prototype credibility to customer-grade repeatability. The next milestone is not to build every integration deeply. It is to create a reliable extension model, a few high-value templates, and measurable customer outcomes."
        ],
        "table": [
            ("0-3 months", "Stabilize workflow UI, template import/export, CellX DB nodes, manual ChatGPT handoff, marketing site, and one-person-company demo flows.", "Competition-ready and pilot-ready prototype."),
            ("3-6 months", "Add production sidecar API, encrypted secrets, connector tests, run logs, and templates for social, CRM, Amazon, FedEx/UPS, Stripe/PayPal, and email.", "Paid pilot readiness."),
            ("6-12 months", "Add marketplace templates, multi-tenant account model, billing, usage controls, and role-based workflow governance across industry packs.", "Repeatable SaaS motion."),
            ("12-18 months", "Partner program, enterprise deployment kit, advanced AI agent orchestration, and analytics.", "Scale beyond founder-led implementation.")
        ],
    },
    {
        "title": "18. Revenue Forecast",
        "body": [
            "The following forecast is an investor-case planning model, not an audited financial forecast. It assumes CellX begins with high-touch paid pilots, converts the most repeatable workflows into monthly subscriptions, and uses partners plus template packs to accelerate adoption. Revenue is split into four layers: subscription, implementation, template and integration packs, and pass-through usage.",
            "The upside case is built around a narrow wedge that can expand quickly after proof. By Year 1, CellX aims to show paid demand and referenceable outcomes. By Year 2, partner and template motions begin to scale monthly recurring revenue. By Year 3, the target is approximately $6.0M ARR run-rate, supported by 900 paid accounts, higher Growth/Enterprise mix, and reusable workflow assets."
        ],
        "custom_table": {
            "headers": ["Investor Case", "Year 1", "Year 2", "Year 3"],
            "rows": [
                ("Ending paid accounts", "35", "220", "900"),
                ("Ending MRR run-rate", "$15k", "$104k", "$500k"),
                ("ARR run-rate", "$180k", "$1.25M", "$6.0M"),
                ("Recognized subscription revenue", "$80k", "$650k", "$3.45M"),
                ("Implementation revenue", "$135k", "$360k", "$720k"),
                ("Templates, connectors, usage", "$35k", "$190k", "$630k"),
                ("Total recognized revenue", "$250k", "$1.20M", "$4.80M")
            ],
        },
        "bullets": [
            "Year 1 account mix assumption: 22 Starter, 10 Growth, 3 Enterprise by year end.",
            "Year 2 account mix assumption: 110 Starter, 85 Growth, 25 Enterprise by year end.",
            "Year 3 account mix assumption: 360 Starter, 400 Growth, 140 Enterprise by year end.",
            "Investor upside depends on repeatable acquisition, partner delivery, and workflow expansion after the first paid use case."
        ],
    },
    {
        "title": "19. Profitability Analysis",
        "body": [
            "CellX can become profitable when three things happen together: subscription revenue grows, implementation work is productized into reusable templates, and AI usage costs are either passed through or owned by the customer. The platform should avoid absorbing unpredictable model costs in low-priced plans. BYOK and manual handoff modes protect gross margin while still letting customers use AI.",
            "The operating model should stay lean through the first three years, but the investor case assumes faster commercial execution and stronger partner leverage than the conservative case. Product development, cloud infrastructure, security review, connector maintenance, and customer success are the main costs. Break-even becomes realistic when recurring revenue and template margin grow faster than founder-led implementation costs."
        ],
        "custom_table": {
            "headers": ["Profit Driver", "Year 1", "Year 2", "Year 3"],
            "rows": [
                ("Gross margin", "68%", "76%", "84%"),
                ("Estimated COGS", "$80k", "$288k", "$768k"),
                ("Operating expense", "$430k", "$1.00M", "$2.78M"),
                ("Operating profit", "-$260k", "-$90k", "$1.25M"),
                ("Break-even logic", "Pilot validation", "Repeatable sales", "SaaS and template leverage")
            ],
        },
        "bullets": [
            "Keep cloud and model costs variable through usage limits, BYOK, pass-through billing, and manual AI handoff.",
            "Use paid pilots to fund discovery while building template assets that can be sold repeatedly.",
            "Prioritize integrations that directly unlock revenue workflows: social distribution, CRM, payments, shipping, Amazon/Shopify, email, and Excel exports.",
            "Break-even target: reach during Year 3 once ARR run-rate approaches $6.0M and implementation delivery becomes partner-assisted."
        ],
    },
    {
        "title": "20. Thank You, Competition Ask, and References",
        "body": [
            "Presented by Harrison Huang and David Cai from Tarbut V' Torah (TVT) Community Day School, 9th Grade.",
            "Competition ask: introductions to pilot customers, cloud credits, partner channels, and mentoring on enterprise AI governance. Funding, if pursued, should prioritize product hardening, connector development, security review, and customer implementation support."
        ],
        "bullets": [
            "Gartner, Forecast Analysis: Low-Code Development Technologies, Worldwide, 2025: https://www.gartner.com/en/documents/7146430",
            "Forrester, The Low-Code Market Could Approach $50 Billion By 2028: https://www.forrester.com/blogs/the-low-code-market-could-approach-50-billion-by-2028/",
            "Deloitte, State of AI in the Enterprise 2026 press release: https://www.deloitte.com/us/en/about/press-room/state-of-ai-report-2026.html",
            "McKinsey, Reimagining tech infrastructure for agentic AI: https://www.mckinsey.com/capabilities/mckinsey-technology/our-insights/reimagining-tech-infrastructure-for-and-with-agentic-ai",
            "Palantir 2025 Form 10-K and public revenue/price data: https://www.sec.gov/Archives/edgar/data/1321655/000132165526000011/pltr-20251231.htm",
            "UiPath FY2026 Form 10-K: https://ir.uipath.com/financials/sec-filings/content/0001734722-26-000012/path-20260131.htm",
            "ServiceNow 2025 results: https://newsroom.servicenow.com/press-releases/details/2026/ServiceNow-Reports-Fourth-Quarter-and-Full-Year-2025-Financial-Results-Board-of-Directors-Authorizes-Additional-5B-for-Share-Repurchase-Program/default.aspx",
            "Salesforce FY2026 investor FAQ and recent Agentforce/Data 360 reporting: https://investor.salesforce.com/resources/investor-faqs/",
            "C3.ai FY2026 results: https://c3.ai/news/c3-ai-announces-fiscal-fourth-quarter-and-full-fiscal-year-2026-results",
            "Microsoft FY2026 Form 10-K: https://www.sec.gov/Archives/edgar/data/789019/000119312526323660/msft-20260630.htm"
        ],
    },
]


def build_docx():
    OUT.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.68)
    section.right_margin = Inches(0.68)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.2)

    # Cover
    add_visual(doc, "cellx_hero_ad.png", 7.1)
    add_para(doc, "Competition Business Plan", 13, True, BLUE, 8)
    add_para(doc, "CellX RDP", 34, True, NAVY, 4)
    add_para(doc, "Rapid Development Platform for Data Apps, AI Workflow Skills, and Enterprise Integrations", 16, False, "334155", 14)
    add_para(doc, "Software-defined software for the AI-agent era: flexible business processes, one-person-company leverage, and cost-down automation.", 12.5, True, BLUE, 12)
    add_para(doc, "Presented by Harrison Huang and David Cai", 11.5, True, NAVY, 3)
    add_para(doc, "Tarbut V' Torah (TVT) Community Day School | 9th Grade", 10.5, False, MUTED, 12)
    add_para(doc, "Prepared for startup and innovation competition review", 10.5, False, MUTED, 18)
    add_metric_table(doc, [
        ("Public site", "https://cellaidata.com", "Lead generation, pricing, demo entry, workflow promotion."),
        ("Backend", "https://app.cellaidata.com", "Operational CellX admin platform and data system."),
        ("Workflow", "https://app.cellaidata.com/workflow/", "AI workflow skill design and extension surface.")
    ])
    add_para(doc, "This document is an English competition draft based on the current CellX prototype, deployed pages, workflow designer work, and public market sources listed in the reference page.", 9.2, False, MUTED, 2)

    for idx, page in enumerate(pages):
        add_heading(doc, page["title"])
        for para in page.get("body", []):
            add_para(doc, para)
        if page["title"].startswith("4."):
            add_visual(doc, "cellx_architecture.png", 6.75)
        if page["title"].startswith("5."):
            add_visual(doc, "cellx_workflow_ad.png", 6.75)
        if page["title"].startswith("8."):
            add_visual(doc, "cellx_market_ad.png", 6.75)
        if page["title"].startswith("15."):
            add_visual(doc, "cellx_architecture.png", 6.75)
        if "table" in page:
            add_metric_table(doc, page["table"])
        if "custom_table" in page:
            add_table(doc, page["custom_table"]["headers"], page["custom_table"]["rows"])
        if "bullets" in page:
            add_bullets(doc, page["bullets"])
        if idx != len(pages) - 1:
            doc.add_page_break()

    doc.save(DOCX)
    print(DOCX.resolve())


if __name__ == "__main__":
    build_docx()
